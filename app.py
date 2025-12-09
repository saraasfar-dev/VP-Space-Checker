import streamlit as st
import docx
import re
from io import BytesIO
import random

# ================================================================
#                     AUTHENTICATION BLOCK
# ================================================================
ALLOWED_DOMAIN = "@softwarefinder.com"

# Session state initialization
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "otp_sent" not in st.session_state:
    st.session_state.otp_sent = False
if "generated_otp" not in st.session_state:
    st.session_state.generated_otp = None
if "email_verified" not in st.session_state:
    st.session_state.email_verified = None

# Step 1: Enter email
def show_email_page():
    st.title("Software Finder – Secure Login")
    st.write("Enter your work email to continue.")
    email = st.text_input("Work Email")
    if st.button("Send Verification Code"):
        if not email.endswith(ALLOWED_DOMAIN):
            st.error(f"Access restricted to {ALLOWED_DOMAIN} users.")
            return
        # Generate a fake OTP for testing (displayed here; no email sent)
        otp = random.randint(1000, 9999)
        st.session_state.generated_otp = otp
        st.session_state.email_verified = email
        st.session_state.otp_sent = True
        st.success(f"Your verification code is: {otp}")
        st.info("In production, OTP would be sent to your email.")

# Step 2: Enter OTP
def show_otp_page():
    st.title("Enter Verification Code")
    user_otp = st.text_input("4-digit code", max_chars=4)
    if st.button("Verify"):
        if user_otp == str(st.session_state.generated_otp):
            st.success("Authentication successful!")
            st.session_state.authenticated = True
        else:
            st.error("Incorrect code. Please try again.")

# Gatekeeper
if not st.session_state.authenticated:
    if not st.session_state.otp_sent:
        show_email_page()
    else:
        show_otp_page()
    st.stop()  # Stop until authentication is done

# ================================================================
#                     MAIN APP STARTS HERE
# ================================================================
# App title and logo
st.set_page_config(page_title="SF VP Spacing & Meta Checker", layout="wide")
st.image("logo.png", width=150)
st.title("SF VP Spacing & Meta Checker")

st.write("This tool auto-fixes spacing issues in Vendor Profiles and checks Meta Titles/Descriptions for character length compliance.")

# --- Function to fix spacing issues ---
def fix_spacing_issues_inplace(doc):
    spacing_issues = []
    for para in doc.paragraphs:
        original_text = para.text
        if re.search(r"#\w+#\s", original_text) or re.search(r"\s#\w+#", original_text):
            spacing_issues.append(original_text.strip())
        corrected_text = re.sub(r"(#\w+#)\s+", r"\1", original_text)
        corrected_text = re.sub(r"\s+(#\w+#)", r"\1", corrected_text)
        if corrected_text != original_text:
            para.text = corrected_text
    return doc, spacing_issues

# --- Function to check meta titles and descriptions ---
def check_meta_limits(doc):
    issues = []
    patterns = {
        "Meta Title": r"#smts#(.*?)#smte#",
        "Meta Description": r"#smds#(.*?)#smde#",
        "Review Meta Title": r"#rmts#(.*?)#rmte#",
        "Review Meta Description": r"#rmds#(.*?)#rmde#",
        "Alternative Meta Title": r"#amts#(.*?)#amte#",
        "Alternative Meta Description": r"#amds#(.*?)#amde#"
    }
    for para in doc.paragraphs:
        text = para.text
        for label, pattern in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                content = match.group(1).strip()
                length = len(content)
                if "Title" in label and length > 80:
                    issues.append(f"<span style='color:red;'>❌ {label} exceeds 80 chars ({length}): {content}</span>")
                elif "Description" in label and (length < 130 or length > 180):
                    issues.append(f"<span style='color:red;'>❌ {label} out of range (130–180). Length {length}: {content}</span>")
    return issues

# --- File uploader ---
uploaded_file = st.file_uploader("Upload a Vendor Profile (.docx)", type=["docx"])
if uploaded_file:
    doc = docx.Document(uploaded_file)
    corrected_doc, spacing_issues = fix_spacing_issues_inplace(doc)
    meta_issues = check_meta_limits(doc)

    st.subheader("Spacing Issues Detected (Auto-Fixed in Download):")
    if spacing_issues:
        for issue in spacing_issues:
            st.write(f"- {issue}")
    else:
        st.success("✅ No spacing issues found!")

    st.subheader("Meta Title & Description Issues:")
    if meta_issues:
        for issue in meta_issues:
            st.markdown(issue, unsafe_allow_html=True)
    else:
        st.success("✅ All meta titles and descriptions within limits!")

    buffer = BytesIO()
    corrected_doc.save(buffer)
    buffer.seek(0)
    st.download_button(
        label="📥 Download Corrected File",
        data=buffer,
        file_name="corrected_vendor_profile.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# --- Footer ---
st.write("---")
st.markdown(
    """
    <div style='text-align:center;'>
        <p>Developed by <b>Sarah Asfar</b>  © 2025 Software Finder</p>
    </div>
    """,
    unsafe_allow_html=True
)
